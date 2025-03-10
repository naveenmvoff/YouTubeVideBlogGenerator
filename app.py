import streamlit as st
from dotenv import load_dotenv 
import os 
import google.generativeai as genai
from youtube_transcript_api import YouTubeTranscriptApi
import re
import io
import time
from functools import lru_cache

# Try to import docx, provide fallback if not available
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("python-docx is not installed. Only text download will be available. Install with: pip install python-docx")

load_dotenv()  # Read .env file and load variables into environment
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))  # Retrieve the key from environment and pass the API Key

# Prompt for summarization
prompt = """You are Youtube video summarizer. You will be taking the transcript text and summarizing the entire video and providing the important Heading("The Video Heading") and Introduction("The whole introduction about the video"), Key Points, Notable Quotes, and Conclusion.   Don't need to Bold Anything in the Providing content."""

# Function to extract video ID from YouTube link
def extract_video_id(youtube_link):
    # Regular expression to capture YouTube video ID from various formats
    pattern = r"(?:v=|\/)([a-zA-Z0-9_-]{11})"
    match = re.search(pattern, youtube_link)
    if match:
        return match.group(1)
    return None

# Add caching for transcript fetching
@lru_cache(maxsize=100)
def cached_fetch_transcript(video_id):
    return YouTubeTranscriptApi.get_transcript(video_id)

def extract_transcript_details(youtube_video_url):
    try:
        video_id = extract_video_id(youtube_video_url)
        if not video_id:
            raise ValueError("Invalid YouTube video URL. Could not extract video ID.")
        
        if 'retry_count' not in st.session_state:
            st.session_state.retry_count = 0
        
        max_retries = 5  # Increased retries
        retry_delay = 2  # Increased delay

        while st.session_state.retry_count < max_retries:
            try:
                transcript_text = cached_fetch_transcript(video_id)
                st.session_state.retry_count = 0
                return " ".join(t["text"] for t in transcript_text)
            except Exception as e:
                st.session_state.retry_count += 1
                if st.session_state.retry_count < max_retries:
                    time.sleep(retry_delay * st.session_state.retry_count)  # Exponential backoff
                    continue
                try:
                    transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
                    for lang in ['en', 'en-US', 'en-GB', 'a.en', 'auto']:
                        try:
                            transcript_text = transcript_list.find_transcript([lang]).fetch()
                            return " ".join(t["text"] for t in transcript_text)
                        except:
                            continue
                    raise ValueError("""Unable to access video captions. Please:
                    1. Clear cache and try again
                    2. Verify the video has captions enabled
                    3. Try a different video
                    4. If problem persists, the service might be temporarily unavailable""")
                except Exception as inner_e:
                    raise ValueError("""Service temporarily unavailable. Please:
                    1. Clear cache
                    2. Wait a few minutes
                    3. Try a different video with captions enabled""")
    except Exception as e:
        raise e

# Function to generate summary using Google Gemini Pro
def generate_gemini_content(transcript_text):
    # model = genai.GenerativeModel("gemini-pro")
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(prompt + transcript_text)
    return response.text

def create_word_document(summary):
    doc = Document()
    doc.add_heading('YouTube Video Summary', 0)
    doc.add_paragraph(summary)
    
    # Save document to byte stream
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

# Streamlit interface
st.set_page_config(
    page_title="YouTube Transcript Summarizer",
    page_icon="📝"
    # Removed wide layout
)

# Initialize session state
if 'processed_videos' not in st.session_state:
    st.session_state.processed_videos = set()
if 'current_summary' not in st.session_state:
    st.session_state.current_summary = None
if 'current_video_id' not in st.session_state:
    st.session_state.current_video_id = None
if 'cache_cleared' not in st.session_state:
    st.session_state.cache_cleared = False

st.title("YouTube Content Summarizer")

# Adjust input and button layout
youtube_link = st.text_input("Enter YouTube Video Link")


# Place buttons side by side using columns
col1, col2 = st.columns([1, 4])
with col1:
    if st.button("Clear Cache", help="Clear cached transcripts"):
        try:
            cached_fetch_transcript.cache_clear()
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.success("Cache cleared!")
            time.sleep(0.5)
            st.rerun()
        except Exception as e:
            st.error(f"Error clearing cache: {e}")

with col2:
    if st.button("Get Content", use_container_width=True):
        if youtube_link:
            try:
                with st.spinner('Fetching video transcript...'):
                    transcript_text = extract_transcript_details(youtube_link)

                if transcript_text:
                    with st.spinner('Generating summary...'):
                        summary = generate_gemini_content(transcript_text)
                        st.session_state.current_summary = summary
                        st.session_state.current_video_id = extract_video_id(youtube_link)

            except Exception as e:
                st.error("Unable to process video")
                st.info("""Please try the following:
                1. Click 'Clear Cache' button
                2. Verify the video has captions/subtitles enabled
                3. Try a different video
                4. If issues persist, wait a few minutes and try again""")
        else:
            st.error("Please provide a valid YouTube link.")

# Display content container (outside the if st.button block)
if st.session_state.current_summary:
    st.markdown("# Blog Content:")
    
    if st.session_state.current_video_id:
        st.image(f"http://img.youtube.com/vi/{st.session_state.current_video_id}/0.jpg", use_column_width=True)
    
    st.write(st.session_state.current_summary)
    
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            label="Download as Text",
            data=st.session_state.current_summary.encode('utf-8'),
            file_name="video_summary.txt",
            mime="text/plain",
            key="text_download"
        )
    
    if DOCX_AVAILABLE:
        with col2:
            word_doc = create_word_document(st.session_state.current_summary)
            st.download_button(
                label="Download as Word Document",
                data=word_doc,
                file_name="video_summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="word_download"
            )

